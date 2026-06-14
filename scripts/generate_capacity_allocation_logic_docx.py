# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "Capacity_Allocation_Logic_CN_V02.docx"
CONTENT_WIDTH_INCHES = 7.0


def set_run_font(run, *, size=10.5, bold=False, color="1F2937", font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, *, size, bold=False, color="1F2937") -> None:
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, margin: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_table_width(table, widths: list[float]) -> None:
    scale = min(1.0, CONTENT_WIDTH_INCHES / sum(widths))
    widths = [width * scale for width in widths]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(dxa(sum(widths))))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(dxa(width)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa(width)))
            tc_w.set(qn("w:type"), "dxa")


def write_cell(cell, text: str, *, bold=False, color="1F2937", size=8.0) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins(cell, 70)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        write_cell(cell, header, bold=True, color="0B2545")
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            write_cell(row.cells[idx], value)
    set_table_width(table, widths)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color="2E74B5" if level <= 2 else "1F4D78")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.18
        run = p.add_run(item)
        set_run_font(run, size=10)


def add_formula(doc: Document, formula: str, explanation: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(formula)
    set_run_font(run, size=9.4, bold=True, color="0B2545", font="Consolas")
    if explanation:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.18)
        p2.paragraph_format.space_after = Pt(5)
        r2 = p2.add_run(explanation)
        set_run_font(r2, size=9.4, color="374151")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF7E6")
    set_cell_margins(cell, 140)
    p = cell.paragraphs[0]
    r1 = p.add_run(title + "：")
    set_run_font(r1, size=9.5, bold=True, color="8A4B00")
    r2 = p.add_run(body)
    set_run_font(r2, size=9.5, color="3B2F1F")
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    set_style_font(doc.styles["Normal"], size=10.5)
    set_style_font(doc.styles["Heading 1"], size=16, bold=True, color="2E74B5")
    set_style_font(doc.styles["Heading 2"], size=13, bold=True, color="2E74B5")
    set_style_font(doc.styles["Heading 3"], size=12, bold=True, color="1F4D78")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Capacity Optimizer 产能分配逻辑说明")
    set_run_font(r, size=22, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    r = sub.add_run("适用版本：v2.2.2 之后 | 范围：ModeA、ModeB、routing、setup、超产能分配逻辑")
    set_run_font(r, size=10.5, color="374151")

    add_note(
        doc,
        "核心结论",
        "工具不是按产品顺序逐个塞产能，而是把同一个月内所有需求、所有可用工作中心路径、setup 占用、未满足和外包选择放进一个优化模型中一起求解。产能竞争按“占用比例”计算，而不是把不同产品吨数直接相加。",
    )

    add_heading(doc, "1. 为什么不能直接按吨相加", 1)
    add_para(doc, "同一个工作中心生产不同产品时，每个产品的月产出能力可能不同。因此，50 吨产品 A 和 50 吨产品 B 对工作中心的占用不一定相同。工具先把每个产品在该工作中心上的分配吨位换算成产能占用比例，再判断工作中心是否超载。")
    add_formula(doc, "capacity_usage(p,w) = allocation(p,w) / monthly_capacity(p,w)", "含义：产品 p 在工作中心 w 上分配的吨位，占用了该工作中心当月多少比例的产能。")
    add_formula(doc, "Σ capacity_usage(p,w) + Σ setup_share(f,plant,w) <= 1", "含义：同一个工作中心 w 在同一个月内，所有产品生产占用比例加上 ProductFamily 级 setup 占用比例，不能超过 100%。f 表示 planner 清单 D 列 ProductFamily。")
    add_para(doc, "例子：6 月工作中心 A 对产品 P1 的月产能是 100 吨，对产品 P2 的月产能是 50 吨。现在 P1 分配 50 吨，P2 也分配 50 吨。表面上看是 50 + 50 = 100 吨，但这并不代表工作中心 A 只占用了 100 吨对应的产能。")
    add_table(
        doc,
        ["产品", "分配吨位 allocation(p,A)", "该产品在 A 的月产能 monthly_capacity(p,A)", "产能占用比例"],
        [
            ["P1", "50 吨", "100 吨/月", "50 / 100 = 0.50"],
            ["P2", "50 吨", "50 吨/月", "50 / 50 = 1.00"],
            ["合计", "100 吨", "不能直接相加", "0.50 + 1.00 = 1.50"],
        ],
        [1.0, 1.55, 2.0, 1.55],
    )
    add_para(doc, "这个例子中，如果只看吨位，使用者可能会误以为 100 吨刚好等于工作中心 A 的产能。但按真实生产速率换算后，P1 已经占用 50% 时间，P2 又占用 100% 时间，合计占用 150%。因此工作中心 A 实际已经超载 50%。这就是工具必须使用产能占用比例，而不是直接把不同产品的吨位相加的原因。")
    add_para(doc, "正例：仍然是 6 月工作中心 A，P1 的月产能是 100 吨，P2 的月产能是 50 吨。如果 P1 分配 50 吨，P2 分配 25 吨，则这两个产品虽然合计只有 75 吨，但按各自生产速率折算后，正好占用工作中心 A 的 100% 产能。")
    add_table(
        doc,
        ["产品", "分配吨位 allocation(p,A)", "该产品在 A 的月产能 monthly_capacity(p,A)", "产能占用比例"],
        [
            ["P1", "50 吨", "100 吨/月", "50 / 100 = 0.50"],
            ["P2", "25 吨", "50 吨/月", "25 / 50 = 0.50"],
            ["合计", "75 吨", "不能直接用 75 吨判断", "0.50 + 0.50 = 1.00"],
        ],
        [1.0, 1.55, 2.0, 1.55],
    )
    add_para(doc, "这个正例说明，工具判断工作中心是否满载时，真正看的不是“分配吨位合计是否等于某个吨数”，而是每个产品各自折算出来的产能占用比例之和。只要比例合计等于 1.00，就表示该工作中心当月正好 100% 被占用；如果再考虑 setup，则还需要把 ProductFamily 级 setup_share(f,plant,w) 一起加进去。")

    add_heading(doc, "2. 符号和变量解释", 1)
    add_table(
        doc,
        ["符号", "中文含义", "详细解释"],
        [
            ["p", "Product，产品", "例如 BSEAL HV3WH 12X600ML TU。公式中的 p 表示某一个具体产品。"],
            ["w", "WorkCenter，工作中心", "例如 CA04 - 4000L JD。公式中的 w 表示某一个具体生产资源。"],
            ["m", "Month，月份", "例如 2026-06。产能、需求和 setup 都按月份计算。"],
            ["d", "Demand node，需求节点", "工具内部的需求颗粒度，通常是 d = (m, p, plant, source_resource)。"],
            ["allocation(p,w)", "产品 p 分配到工作中心 w 的吨位", "这是优化器决定的变量，也就是最终报告中的 allocated tons。"],
            ["monthly_capacity(p,w)", "产品 p 在工作中心 w 的月产能", "由 Annual Max Capacity Tons 或 Annual Planned Capacity Tons 除以 12 得到。不同产品在同一工作中心上的值可以不同。"],
            ["f", "ProductFamily，产品族", "来自 planner 清单 D 列。setup 触发按 f + plant + w 聚合。"],
            ["setup_share(f,plant,w)", "ProductFamily f 在某工厂某工作中心的 setup 占用比例", "等于该组 setup_hours / month_hours(m)。只要该 ProductFamily 在该工厂该工作中心有任何内部分配吨位，就触发一次 setup。"],
            ["unmet(d)", "需求节点 d 未满足吨位", "产能不够且无可用替代路径时留下的缺口。"],
            ["outsource(d)", "需求节点 d 外包吨位", "只有 routing 中存在 EligibleFlag = Y 且 Router Type = Toller 的产品才允许使用。"],
        ],
        [1.35, 1.75, 3.4],
    )
    add_para(doc, "例子：如果 2026 年 6 月，产品 P1 在上海工厂的原始来源工作中心是 A，需求 80 吨，并且优化器最终把 60 吨排到 A、20 吨排到 B，那么可以把符号读成下面这样。")
    add_table(
        doc,
        ["符号写法", "例子中的值", "如何理解"],
        [
            ["p", "P1", "正在讨论的具体产品。"],
            ["w", "A 或 B", "产品可以被分配到的具体工作中心。"],
            ["m", "2026-06", "当前计算月份。"],
            ["d", "(2026-06, P1, 上海工厂, A)", "这一条需求节点，来源资源是 A。"],
            ["allocation(P1,A)", "60 吨", "P1 最终排到 A 的吨位。"],
            ["allocation(P1,B)", "20 吨", "P1 最终排到 B 的吨位。"],
        ],
        [1.55, 1.75, 3.3],
    )

    add_heading(doc, "3. 输入如何变成可分配路径", 1)
    add_para(doc, "工具先读取 planner load 形成需求，再读取 capacity 和 routing 形成可用路径。每个需求节点只会在符合条件的工作中心集合中分配。")
    add_table(
        doc,
        ["来源", "决定什么", "参与逻辑"],
        [
            ["planner load", "需求 d 和 demand(d)", "按 Month + Product + Plant + Source_Resource 聚合需求吨位。"],
            ["master_capacity", "主产能路径 Capacity_Base", "如果产品 p 在 source_resource 上有月产能，则该路径可用。"],
            ["master_routing", "Routing_Reroute 和 Toller 路径", "EligibleFalg/EligibleFlag = Y 的非 Toller 路径可作为 alternative；EligibleFlag = Y 且 Router Type = Toller 的产品可外包。"],
            ["Setup_Hours", "setup 占用", "同月同工厂同工作中心同 ProductFamily 有任何内部分配吨位时触发一次 setup。"],
        ],
        [1.55, 2.15, 2.8],
    )
    add_para(doc, "例子：planner load 中有一条 P1 在 A 上的 80 吨需求。capacity 表中 P1-A 有产能，routing 表中 P1-B 是 EligibleFlag=Y 的 Alternative，P1-Toller 是 EligibleFlag=Y 且 Router Type=Toller。工具会形成下面的可选路径集合。")
    add_table(
        doc,
        ["路径", "来自哪张表", "是否可用", "用途"],
        [
            ["P1 -> A", "capacity", "可用", "Capacity_Base，优先使用原始来源工作中心。"],
            ["P1 -> B", "routing", "可用", "Alternative，主路径不足时可转移。"],
            ["P1 -> Toller", "routing", "可用", "外包路径，内部产能不足时可用。"],
            ["P1 -> C", "无 capacity/routing", "不可用", "即使 C 有空闲，P1 没有合法路径也不能分配过去。"],
        ],
        [1.25, 1.35, 1.15, 3.25],
    )

    add_heading(doc, "4. 需求平衡公式", 1)
    add_para(doc, "每个需求节点的需求必须被内部工作中心承接、外包承接，或者留下未满足。")
    add_formula(doc, "Σ x(d,w) + outsource(d) + unmet(d) = demand(d)", "x(d,w) 表示需求节点 d 分配到工作中心 w 的吨位。ModeA 通常没有 outsource(d)。ModeB 中只有有 eligible Toller 的产品才有 outsource(d)。")
    add_bullets(
        doc,
        [
            "如果有足够可用产能，unmet(d) 会尽量为 0。",
            "如果主路径不够，ModeB 会尝试 eligible alternative routing。",
            "如果仍不够且产品允许 Toller，则可以进入 outsource(d)。",
            "如果没有任何可行路径或所有路径都满了，剩余吨位进入 unmet(d)。",
        ],
    )
    add_para(doc, "例子：某需求节点 d 的 demand(d)=120 吨。优化后 A 承接 70 吨，B 承接 30 吨，Toller 外包 15 吨，仍有 5 吨未满足。代入需求平衡公式后如下。")
    add_formula(doc, "70 + 30 + 15 + 5 = 120", "也就是 Σ x(d,w)=100，outsource(d)=15，unmet(d)=5。无论内部、外包还是未满足，四部分合计必须等于原始需求 120 吨。")

    add_heading(doc, "5. 工作中心产能约束", 1)
    add_para(doc, "工作中心约束是产能分配的核心。它把不同产品的吨位都换算为该工作中心的占用比例。")
    add_formula(doc, "Σ [x(d,w) / monthly_capacity(p,w)] + Σ [y(f,plant,w) * setup_hours(f,plant,w) / month_hours(m)] <= 1", "其中 d 中包含产品 p 和 ProductFamily f。y(f,plant,w) 是 0/1 变量，表示 ProductFamily f 在该工厂该工作中心当月是否触发 setup。")
    add_table(
        doc,
        ["公式部分", "解释"],
        [
            ["x(d,w) / monthly_capacity(p,w)", "产品 p 在工作中心 w 上分配 x 吨，占用该工作中心多少比例的月产能。"],
            ["y(f,plant,w)", "是否触发 setup。只要该 ProductFamily 在该工厂该工作中心有任何内部分配，优化器需要把 y 设为 1。"],
            ["setup_hours(f,plant,w) / month_hours(m)", "setup 小时数折算为当月工作中心产能占比。例如 7.2 小时 / 720 小时 = 1%。"],
            ["<= 1", "同一工作中心同一月份总占用不能超过 100%。"],
        ],
        [2.2, 4.1],
    )
    add_para(doc, "例子：同一个月内，工厂 PLT1 的工作中心 A 生产 P1 和 P2，且二者都属于 planner 清单 D 列 ProductFamily F1。P1 在 A 上分配 40 吨，月产能 100 吨；P2 在 A 上分配 20 吨，月产能 50 吨；F1 在 PLT1+A 上只触发一次 setup，setup_share 为 0.01。")
    add_formula(doc, "40/100 + 20/50 + 0.01 = 0.81 <= 1", "说明 A 的总占用是 81%，还有 19% 的产能空间。这里不能用 40+20=60 吨判断是否满载，必须用占用比例判断；同 ProductFamily 的两个产品不会重复计算两次 setup。")

    add_heading(doc, "6. Setup 触发逻辑", 1)
    add_formula(doc, "if Σ x(d,w for ProductFamily f, plant, w) > 0, then y(f,plant,w) = 1", "不再有 1 吨以下不触发的规则。同月、同工厂、同工作中心、同 ProductFamily 只触发一次 setup。")
    add_formula(doc, "setup_equivalent_tons_by_max(f,plant,w) = setup_hours(f,plant,w) * monthly_max_capacity_reference / month_hours(m)", "报告中展示的 setup 等效吨位按 Max 产能换算，即使 Planned 口径下也按 Max 速率换算。同组多个产品存在不同 Setup_Hours 时，模型使用该组中被分配产品的最大 setup 值避免低估。")
    add_para(doc, "需要注意，模型约束里 setup 以时间占比进入工作中心约束；报告里为了让业务用户理解，会额外展示等效吨位。")
    add_para(doc, "直观例子：2026 年 6 月，工厂 PLT1 的工作中心 A 需要生产 3 个产品。P1 和 P2 属于同一个 ProductFamily F1，可以连续生产；P3 属于另一个 ProductFamily F2，需要另一套 setup。")
    add_table(
        doc,
        ["产品", "ProductFamily", "分配到 PLT1 + A 的吨位", "是否单独触发 setup"],
        [
            ["P1", "F1", "0.8 吨", "不单独触发；归入 F1 这一次 setup"],
            ["P2", "F1", "30 吨", "不单独触发；与 P1 共用 F1 setup"],
            ["P3", "F2", "20 吨", "触发 F2 这一次 setup"],
        ],
        [0.8, 1.35, 1.55, 2.3],
    )
    add_para(doc, "因此，虽然有 3 个产品，但 setup 不是按产品数计算，而是按 ProductFamily 在同一工厂同一工作中心上的组数计算。")
    add_table(
        doc,
        ["Setup 组", "组内产品", "该组总分配", "setup 次数", "原因"],
        [
            ["F1 + PLT1 + A", "P1 + P2", "0.8 + 30 = 30.8 吨", "1 次", "同 ProductFamily，可连续生产"],
            ["F2 + PLT1 + A", "P3", "20 吨", "1 次", "不同 ProductFamily，需要另一套 setup"],
            ["合计", "P1 + P2 + P3", "50.8 吨", "2 次", "不是 3 次；因为 P1 和 P2 共用一次"],
        ],
        [1.35, 1.3, 1.45, 0.9, 2.0],
    )
    add_para(doc, "如果 6 月有 30 天，month_hours(m)=720，并且 F1 和 F2 在 PLT1+A 上的 setup 都是 7.2 小时，则每个 setup_share 都是 7.2 / 720 = 0.01。这个例子中 F1 触发一次、F2 触发一次，所以 setup 总占用比例是 0.01 + 0.01 = 0.02。注意 P1 只有 0.8 吨也会触发 F1 这组 setup，因为当前规则已经取消了 1 吨以下不触发的例外。")

    add_heading(doc, "7. 目标函数：超产能时如何决定谁排入", 1)
    add_para(doc, "当所有需求合计超过可用产能时，优化器不是人工指定某个产品优先，而是最小化一组惩罚。惩罚越大，模型越不愿意选择。")
    add_formula(doc, "Minimize BIG_M * unmet + TOLLER_PENALTY * outsource + setup_penalty(f,plant,w) * y + route_penalty * x", "这是简化表达。实际模型会对每个需求节点、每条路径、每个 setup 变量分别累加。")
    add_formula(doc, "setup_penalty(f,plant,w) = SETUP_TRIGGER_PENALTY + setup_hours(f,plant,w) * SETUP_HOURS_PENALTY + setup_equivalent_tons_by_max(f,plant,w) * SETUP_TONS_PENALTY", "含义：ModeB 会先尽量少触发 setup 组；setup 组数相同时，会倾向选择 setup 时间更短、等效吨位更低的路径。")
    add_table(
        doc,
        ["惩罚项", "当前量级", "业务含义"],
        [
            ["unmet", "1,000,000,000 / 吨", "最大惩罚。只要有可行产能，模型会优先减少未满足吨位。"],
            ["setup trigger", "1,000,000 / ProductFamily-工厂-工作中心", "抑制不必要的 ProductFamily 拆分和额外换模。"],
            ["setup hours", "1,000 / 小时", "同样 setup 组数下，ModeB 倾向选择 setup 时间更短的路径。"],
            ["outsource / Toller", "100,000 / 吨", "比 unmet 好，但比内部产能差。"],
            ["Capacity Base route", "最低", "优先使用需求原始来源资源。"],
            ["Routing Primary", "较低", "可用 routing 中优先级较高的路径。"],
            ["Routing Alternative", "较高", "主路径不足时再使用。"],
            ["Priority / PenaltyWeight", "附加微调", "Priority 越小越优先；PenaltyWeight 可以覆盖默认路线惩罚。"],
        ],
        [1.7, 1.65, 3.0],
    )
    add_note(doc, "解释", "因为 unmet 每吨惩罚极大，所以超产能时工具首先考虑怎样满足更多吨位。若两个产品抢同一工作中心，在没有其他业务优先级字段时，产出效率更高的产品通常更容易被排入，因为同样 100% 工作中心占用可以满足更多吨位。")
    add_para(doc, "例子：某月剩余内部产能只能再满足 30 吨，但有 50 吨需求缺口。如果不使用 Toller，会有 50 吨 unmet；如果使用 Toller，可以外包 20 吨，只剩 30 吨 unmet。虽然 Toller 有惩罚，但 unmet 惩罚更高，所以模型会选择使用 Toller 来减少 unmet。")
    add_table(
        doc,
        ["方案", "unmet", "Toller", "简化惩罚比较", "模型倾向"],
        [
            ["不用 Toller", "50 吨", "0 吨", "50 * BIG_M", "较差"],
            ["使用 Toller 20 吨", "30 吨", "20 吨", "30 * BIG_M + 20 * TOLLER_PENALTY", "较优"],
        ],
        [1.3, 1.0, 1.0, 2.45, 1.05],
    )

    add_heading(doc, "8. ModeA 和 ModeB 的区别", 1)
    add_table(
        doc,
        ["模式", "可用路径", "超产能处理"],
        [
            ["ModeA", "主要使用 capacity 中需求来源工作中心对应的产能。", "主路径不够时，无法被主路径承接的部分进入 unmet。"],
            ["ModeB", "把 Capacity_Base、eligible routing alternative、eligible Toller 和 unmet 放在一个全局模型中一起求解。", "先尽量用内部可行路径满足需求，再根据惩罚使用 Toller，最后才留下 unmet。"],
        ],
        [1.1, 2.65, 2.75],
    )
    add_para(doc, "例子：P1 的 planner 来源工作中心是 A，需求 100 吨。A 只能承接 70 吨；routing 表里 P1-B 是 eligible alternative，B 还能承接 25 吨。")
    add_table(
        doc,
        ["模式", "可用分配", "结果解释"],
        [
            ["ModeA", "A=70，unmet=30", "只看主来源能力，B 不参与承接。"],
            ["ModeB", "A=70，B=25，unmet=5", "eligible alternative B 可参与，未满足吨位从 30 降到 5。"],
        ],
        [1.1, 2.0, 3.3],
    )

    add_heading(doc, "9. 例子一：同一工作中心，不同产品产能不同", 1)
    add_table(
        doc,
        ["产品", "工作中心 A 月产能", "需求", "说明"],
        [
            ["P1", "100 吨/月", "100 吨", "P1 在 A 上产出效率较高。"],
            ["P2", "50 吨/月", "50 吨", "P2 在 A 上产出效率较低。"],
        ],
        [1.0, 1.7, 1.1, 2.7],
    )
    add_formula(doc, "x(P1,A)/100 + x(P2,A)/50 <= 1", "工作中心 A 的产能约束。")
    add_para(doc, "如果都全量生产：")
    add_formula(doc, "100/100 + 50/50 = 2 > 1", "说明 A 不可能同时满足 P1 100 吨和 P2 50 吨。")
    add_para(doc, "比较几个可能方案：")
    add_table(
        doc,
        ["方案", "分配", "占用计算", "满足吨位", "未满足吨位"],
        [
            ["只排 P1", "P1=100, P2=0", "100/100 = 1", "100", "50"],
            ["只排 P2", "P1=0, P2=50", "50/50 = 1", "50", "100"],
            ["各排一半", "P1=50, P2=25", "50/100 + 25/50 = 1", "75", "75"],
        ],
        [1.0, 1.55, 2.25, 0.9, 0.9],
    )
    add_para(doc, "由于模型首先最小化 unmet，方案“只排 P1”满足 100 吨，未满足 50 吨，比其他方案的 unmet 更少。因此在没有其他业务优先级或 alternative 的情况下，P1 会优先排入 A。")

    add_heading(doc, "10. 例子二：加入 setup 后", 1)
    add_para(doc, "假设 6 月有 30 天，即 month_hours = 30 * 24 = 720 小时。P1 和 P2 都属于 ProductFamily F1，且 F1 在工厂 PLT1 的工作中心 A 上 setup 是 7.2 小时。")
    add_formula(doc, "setup_share(F1,PLT1,A) = 7.2 / 720 = 0.01")
    add_para(doc, "如果只生产 P1，并且 P1 所属 ProductFamily F1 在 A 上有任何内部分配，则约束为：")
    add_formula(doc, "x(P1,A)/100 + 0.01 <= 1")
    add_formula(doc, "x(P1,A) <= 99")
    add_para(doc, "也就是说，setup 会占用 1% 的工作中心月产能，P1 最多只能生产 99 吨。如果 P1 和 P2 同属 ProductFamily F1，并且都在 PLT1+A 上生产，则只触发一次 F1 的 setup：")
    add_formula(doc, "x(P1,A)/100 + x(P2,A)/50 + 0.01 <= 1")
    add_para(doc, "这会让同一个 ProductFamily 不必要地拆到多个工作中心变得更不经济。模型会倾向减少不必要的 ProductFamily-工厂-工作中心组合，因为每多一个组合，都可能触发一次 setup。")

    add_heading(doc, "11. 例子三：有 Alternative Routing 时", 1)
    add_table(
        doc,
        ["产品", "路径", "月产能", "Route Type", "EligibleFalg"],
        [
            ["P1", "A", "100", "Capacity Base", "Y"],
            ["P2", "A", "50", "Capacity Base", "Y"],
            ["P2", "B", "40", "Alternative", "Y"],
        ],
        [0.8, 0.9, 1.0, 1.6, 1.2],
    )
    add_para(doc, "需求仍为 P1=100、P2=50。A 无法同时承接两者。ModeB 会把 P2 去 B 的 alternative 放进可选路径。")
    add_formula(doc, "A constraint: x(P1,A)/100 + x(P2,A)/50 + setup <= 1")
    add_formula(doc, "B constraint: x(P2,B)/40 + setup <= 1")
    add_para(doc, "可行性比较：")
    add_table(
        doc,
        ["方案", "结果", "业务含义"],
        [
            ["P1 留在 A，P2 去 B", "P1 可接近满量；P2 在 B 上最多约 40 吨；P2 可能剩余约 10 吨 unmet。", "总 unmet 约 10，较优。"],
            ["P2 留在 A，P1 不排", "P2 可接近 50 吨；P1 约 100 吨 unmet。", "总 unmet 约 100，明显更差。"],
        ],
        [1.4, 3.2, 1.7],
    )
    add_para(doc, "因为 unmet 惩罚最高，模型会选择 unmet 更少的方案，即 P1 排入 A，P2 尽量排入 B。Alternative 的路线惩罚虽然高于 Capacity Base，但远低于 unmet，所以只要 alternative 能减少 unmet，它就会被使用。")

    add_heading(doc, "12. 当前模型没有做什么", 1)
    add_bullets(
        doc,
        [
            "没有按 planner 或客户重要性自动排序，除非未来新增业务优先级字段。",
            "没有按产品在输入文件中的行顺序决定谁先排。",
            "没有把不同产品的吨数直接相加判断工作中心是否满载。",
            "没有让 Utilization Target 参与当前产能计算，该字段目前只是兼容字段。",
            "没有把 setup 重复算到同月同工厂同工作中心同 ProductFamily 的每一行，只触发一次。",
        ],
    )
    add_para(doc, "例子：P1 是普通订单，P2 是业务上更紧急的订单，但输入表里没有“业务优先级”字段。如果 P1 在 A 上 100% 占用可以满足 100 吨，而 P2 在 A 上 100% 占用只能满足 40 吨，模型可能优先排 P1，因为这样 unmet 更少。这不是模型知道 P1 更重要，而是当前输入没有告诉模型 P2 更重要。")
    add_table(
        doc,
        ["情况", "当前模型会怎么理解", "如果业务希望改变结果"],
        [
            ["P2 更紧急，但没有优先级输入", "模型只看到吨位、路径、产能、setup 和惩罚。", "需要新增业务优先级或 penalty 权重输入。"],
            ["输入文件中 P2 排在 P1 前面", "行顺序不等于生产优先级。", "不能依靠 Excel/CSV 行顺序控制排产。"],
        ],
        [1.6, 2.65, 2.05],
    )

    add_heading(doc, "13. 判断结果是否合理的检查方法", 1)
    add_bullets(
        doc,
        [
            "先看 Allocation_Detail 是否按 Capacity_Basis、Month、Plant、WorkCenter、Product 排序，相关行应聚在一起。",
            "对某个工作中心，手工计算每个产品 allocated_tons / monthly_capacity，再加 ProductFamily 级 setup_share，总和应不超过 1。",
            "如果某产品没有排入，检查它是否有 eligible routing、是否有该产品在该工作中心的 capacity 行、是否被 setup 后产能挤出。",
            "如果更高吨产出的产品被优先满足，通常是因为它在同样产能占用下能减少更多 unmet。",
            "如果业务希望低效率但高优先级的产品先排，需要新增业务优先级权重，否则当前模型不会自动知道这个业务偏好。",
        ],
    )
    add_para(doc, "例子：检查 2026 年 6 月工厂 PLT1 工作中心 A 的 Allocation_Detail，看到 P1 分配 40 吨、P2 分配 20 吨，二者同属 ProductFamily F1；P1-A 月产能 100 吨，P2-A 月产能 50 吨，F1 在 PLT1+A 上只触发一次 setup，setup_share 为 0.01。")
    add_formula(doc, "40/100 + 20/50 + 0.01 = 0.81", "因为 0.81 <= 1，所以这组分配在工作中心 A 上是合理的。如果同样手工算出来大于 1，就需要检查报告口径、月份、工厂、工作中心、ProductFamily、产品产能或 setup 是否对应错了。")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Capacity Optimizer 产能分配逻辑说明")
    set_run_font(r, size=8, color="6B7280")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
