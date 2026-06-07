# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "Capacity_Optimizer_v2.2.1_User_Guide_CN.docx"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, *, size: float, bold: bool = False, color: str = "1F2937") -> None:
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, margin: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def write_cell(cell, text: str, *, bold: bool = False, color: str = "1F2937", size: float = 8.5) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "EAF2F8")
        write_cell(cell, header, bold=True, color="0B2545", size=8.5)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            write_cell(row.cells[idx], value)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=15 if level == 1 else 12, bold=True, color="0B2545" if level == 1 else "1F4E79")


def add_para(doc: Document, text: str, *, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=9.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=9.2)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=9.2)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [6.25])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF7E6")
    set_cell_margins(cell, 140)
    p = cell.paragraphs[0]
    r1 = p.add_run(title + "：")
    set_run_font(r1, size=9.2, bold=True, color="8A4B00")
    r2 = p.add_run(body)
    set_run_font(r2, size=9.2, color="3B2F1F")
    doc.add_paragraph()


def build_document() -> None:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    set_style_font(doc.styles["Normal"], size=9.5)
    set_style_font(doc.styles["Heading 1"], size=15, bold=True, color="0B2545")
    set_style_font(doc.styles["Heading 2"], size=12, bold=True, color="1F4E79")
    set_style_font(doc.styles["Heading 3"], size=10.5, bold=True, color="374151")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Chemical Capacity Optimizer 操作指引")
    set_run_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("版本：v2.2.1 | 更新日期：2026-06-07 | 范围：桌面 GUI、ModeA、ModeB、Max/Planned 口径、Setup Time 计算与报告")
    set_run_font(r, size=10.5, color="374151")

    add_note(
        doc,
        "使用目的",
        "本文用于说明工具如何使用、每个参数如何设置、输入数据如何维护、优化逻辑如何计算、报告如何排序，以及报告中主要列的含义和背后计算逻辑。",
    )

    add_heading(doc, "1. 快速开始", 1)
    add_numbered(
        doc,
        [
            "打开打包后的 dist/CapacityOptimizer/CapacityOptimizer.exe，或在源码模式双击 CapacityOptimizerLauncher.pyw。",
            "首次使用时在启动器中点击 Initialize Workspace，生成用户工作区、输入目录、输出目录、日志目录和授权目录。",
            "维护 Data_Input 下的 planner3_load.csv、master_capacity.csv、master_routing.csv。当前活动输入只保留 planner3_load.csv，其余历史 planner 文件在 Data_Input/archive 中归档。",
            "确认 capacity 与 routing 中同一 Product + Resource 的 Max、Planned、Setup_Hours 完全一致；允许 0，但不允许空。",
            "在启动器中设置 Workspace、Run Mode、Output File Name、Scenario、Start Year、Start Month、Horizon Months、Verbose、Skip Validation Errors 等参数，然后点击 Save Settings。",
            "点击 Run Optimizer。运行结束后点击 Open Output Folder 或到启动器配置的输出目录中查看 ModeA、ModeB、校验报告和明细报告。",
        ],
    )

    add_heading(doc, "2. 目录和文件", 1)
    add_table(
        doc,
        ["路径", "用途", "维护建议"],
        [
            ["CapacityOptimizer.exe / CapacityOptimizerLauncher.pyw", "当前主 UI。用户在桌面 GUI 中初始化工作区、设置参数、生成机器码、运行优化、打开输出和日志。", "业务用户应使用这个入口，不再打开 Excel 控制面板运行工具。"],
            ["launcher_settings.json", "启动器设置文件。保存 GUI 中填写的工作区、输出文件名、运行模式、时间窗口等参数。", "正常情况下由启动器写入，不建议手工编辑。"],
            ["Data_Input/planner3_load.csv", "需求输入。包含月份、planner、产品、工厂、需求吨位等信息。", "这是当前活动 planner load 文件。"],
            ["Data_Input/master_capacity.csv", "主产能表。定义产品在工作中心上的 Max/Planned 产能，以及 Setup_Hours。", "产品主要可生产工作中心应维护在这里。"],
            ["Data_Input/master_routing.csv", "可选 routing 表。定义产品可转移到的替代工作中心，以及同口径产能和 Setup_Hours。", "作为 alternative routing 使用，不要求完全覆盖 capacity。"],
            ["Data_Input/archive", "历史输入归档。包含旧 planner load 和旧版本 master 文件。", "用于追溯，不作为默认运行输入。"],
            ["output", "默认输出目录。保存运行日志、校验结果、ModeA/ModeB 报告。", "每次运行按时间戳或模式生成结果文件。"],
            ["Archive/legacy_excel_control_panel/Capacity_Optimizer_Control.xlsx", "已归档的历史 Excel 控制簿。旧 CLI workbook 模式和部分回归测试仍可使用。", "当前业务用户不应把它当作 UI；正式操作以桌面启动器为准。"],
            ["dist/CapacityOptimizer", "PyInstaller 打包后的 one-folder 工具。", "发布给使用者时优先使用该目录或对应 zip，其中入口是 CapacityOptimizer.exe。"],
        ],
        [1.65, 2.9, 1.75],
    )

    add_heading(doc, "3. 启动器参数", 1)
    add_table(
        doc,
        ["参数", "含义", "建议设置", "错误影响"],
        [
            ["Workspace Root", "用户工作区根目录。启动器在这里维护 Data_Input、output、logs、licenses 等目录。", "打包运行时通常使用默认工作区；源码运行时可指向当前项目工作区。", "路径错误会导致输入、输出或授权文件找不到。"],
            ["Output File Name", "输出报告基础文件名。工具会按模式和时间戳生成实际文件。", "使用清晰名称，例如 capacity_result.xlsx。", "名称不规范时工具会自动补 .xlsx，但不建议使用特殊字符。"],
            ["Scenario Name", "需要运行的场景或全部场景。", "不确定时选择 All。", "设置错误可能导致读取不到期望场景。"],
            ["Run Mode", "运行模式。通常支持 ModeA、ModeB 或 Both。", "需要比较结果时选择 Both；只看全局路径优化时可选 ModeB。", "模式配置不合法会在运行前失败。"],
            ["Start Year / Start Month", "分析起始年月。", "与 planner load 的月份范围保持一致。", "超出输入数据范围会导致有效需求为空或报告不完整。"],
            ["Horizon Months", "分析跨度月份数。", "常用 12；短期分析可按业务窗口缩短。", "跨度过小会漏掉需求，跨度过大会增加计算和报告体量。"],
            ["Verbose", "是否输出更详细日志。", "排查问题时选择 Yes；日常运行可选择 No。", "不影响计算，只影响日志详细程度。"],
            ["Skip Validation Errors", "是否跳过部分校验错误。", "正式运行建议 No。setup 为空、不一致等关键错误不应跳过。", "设置为 Yes 可能掩盖数据维护问题，不建议用于正式结果。"],
        ],
        [1.4, 2.0, 2.15, 2.1],
    )
    add_note(
        doc,
        "旧 Excel 控制簿说明",
        "Archive/legacy_excel_control_panel/Capacity_Optimizer_Control.xlsx 是历史兼容资产。当前 UI 已迁移到 PySide6 桌面启动器，业务用户不需要打开该 Excel 文件设置参数或点击运行。",
    )

    add_heading(doc, "4. 输入表维护", 1)
    add_heading(doc, "4.1 Planner Load", 2)
    add_para(doc, "Planner Load 是需求侧输入。工具会按月份、工厂、产品、planner 等维度读取需求，并在优化中尝试分配到可用工作中心。")
    add_table(
        doc,
        ["列", "含义", "计算用途"],
        [
            ["Month", "需求月份。A 列月份还用于确定自然月天数。", "月小时数 = 当月自然天数 × 24，用于 setup 小时换算成吨。"],
            ["PlannerName", "需求所属 planner。", "用于报告归属、汇总和排序。"],
            ["Product", "产品代码。", "与 capacity/routing 的 Product 匹配。"],
            ["Plant", "工厂。", "用于工厂维度汇总和明细排序。"],
            ["Demand_Tons", "当月需求吨位。", "优化目标是尽量满足该吨位。"],
            ["Customer/Case 字段", "客户或案例维度字段，具体以输入表列名为准。", "ModeB 客户案例报告会继承这些维度。"],
        ],
        [1.35, 2.4, 2.05],
    )

    add_heading(doc, "4.2 Master Capacity", 2)
    add_para(doc, "Master Capacity 是产品主要产能来源表。它表达“这个产品本来应该在哪些工作中心生产，以及这些工作中心每月可提供多少产能”。")
    add_table(
        doc,
        ["列", "是否必填", "含义", "维护规则"],
        [
            ["Product", "是", "产品代码。", "必须与 planner load 的 Product 口径一致。"],
            ["Resource", "是", "工作中心。", "同一个工作中心名称必须全表一致，避免空格和大小写差异。"],
            ["Max Capacity", "是", "最大产能吨位。", "用于 Max 口径计算，也用于 setup 小时换算吨位。"],
            ["Planned Capacity", "是", "计划产能吨位。", "用于 Planned 口径计算；不用于 setup 换算。"],
            ["Setup_Hours", "是", "该产品在该工作中心每次切换/换模需要的小时数。", "允许 0；不允许空；必须是数字；不能为负数。"],
        ],
        [1.15, 0.8, 2.2, 2.3],
    )

    add_heading(doc, "4.3 Master Routing", 2)
    add_para(doc, "Master Routing 是 alternative routing 表。它表达“当产品需要转移时，还可以去哪些替代工作中心生产”。它不是 capacity 的重复表，因此不要求被 capacity 完全覆盖。")
    add_table(
        doc,
        ["列", "是否必填", "含义", "维护规则"],
        [
            ["Product", "是", "产品代码。", "与 capacity 同一 Product + Resource 重叠时必须一致。"],
            ["Resource", "是", "替代工作中心。", "可以出现 capacity 没有覆盖的替代工作中心。"],
            ["Max Capacity", "是", "替代工作中心最大产能。", "若与 capacity 同 Product + Resource 重叠，必须一致。"],
            ["Planned Capacity", "是", "替代工作中心计划产能。", "若与 capacity 同 Product + Resource 重叠，必须一致。"],
            ["Setup_Hours", "是", "该产品在该替代工作中心的换模小时数。", "若与 capacity 同 Product + Resource 重叠，必须一致。"],
        ],
        [1.15, 0.8, 2.2, 2.3],
    )

    add_heading(doc, "5. 输入校验逻辑", 1)
    add_bullets(
        doc,
        [
            "Setup_Hours 在 capacity 和 routing 中均为必填。空值会被视为维护错误，工具停止计算。",
            "Setup_Hours 允许为 0。含义是该产品在该工作中心切换时不消耗额外 setup 时间。",
            "Setup_Hours 必须为数字，不能为负数。",
            "同一个 Product + Resource 若同时出现在 capacity 和 routing 中，则 Max Capacity、Planned Capacity、Setup_Hours 必须一致。",
            "数值一致性容忍度为小数点后四位，即 0.0001。超过该差异会停止计算并在错误报告中列出问题位置。",
            "如果 routing 中出现 capacity 没有的 Product + Resource，不自动报错；这是允许的 alternative routing 场景。",
        ],
    )

    add_heading(doc, "6. Setup Time 计算逻辑", 1)
    add_para(doc, "Setup time 代表同一工作中心从生产一个产品切换到生产另一个产品前需要消耗的准备时间。工具按“同月、同工作中心、同产品”聚合判断是否触发 setup。")
    add_table(
        doc,
        ["规则", "说明"],
        [
            ["触发条件", "当同一个 Month + WorkCenter + Product 的累计分配吨位大于 1 吨时，触发一次 setup。小于或等于 1 吨不触发。"],
            ["同月只算一次", "同一月份、同一工作中心、同一产品即使来自多个 planner 或多条需求，只计算一次 setup。实际生产中通常会把同产品集中排产，因此该逻辑避免重复扣减。"],
            ["月小时数", "Month 对应自然月天数 × 24。例如 6 月为 30 × 24 = 720 小时。"],
            ["吨位换算", "Setup_Equivalent_Tons_By_Max = Setup_Hours × 当月 Max 吨/小时。"],
            ["当月 Max 吨/小时", "当月 Max 吨/小时 = 该产品在该工作中心的月度 Max Capacity ÷ 月小时数。"],
            ["Planned 口径", "即使报告在 Planned 口径下，setup 等效吨仍按 Max 产能换算，不按 Planned 换算。"],
            ["产能占用", "Capacity_Used_Tons = Allocation_Tons + Setup_Equivalent_Tons_By_Max。"],
        ],
        [1.8, 4.45],
    )
    add_note(
        doc,
        "例子",
        "6 月工作中心 A 生产 aa、bb、cc。aa 分配 50 吨、bb 分配 30 吨、cc 分配 20 吨。原逻辑只看 100 吨产能是否够；v2.2.1 后，每个产品若分配吨位大于 1 吨，还会分别加上 aa、bb、cc 对应的 setup 等效吨。",
    )

    add_heading(doc, "7. 优化逻辑", 1)
    add_heading(doc, "7.1 ModeA", 2)
    add_para(doc, "ModeA 以产品需求和工作中心可用产能为核心，生成 Max 与 Planned 两套口径结果。它会优先使用产品的主产能路径，并在需要时结合 routing 可选路径进行分配。")
    add_bullets(
        doc,
        [
            "需求来源：planner load。",
            "主路径来源：master_capacity。",
            "替代路径来源：master_routing。",
            "约束：每个工作中心在每个月、每个产能口径下的 Capacity_Used_Tons 不能超过可用产能。",
            "Setup 作为额外产能占用进入优化模型，而不是报告后手工追加。",
            "当 setup 导致产能不足时，优化会在可行路径中重新分配；若仍不可行，则体现为未满足需求或压力报告。",
        ],
    )

    add_heading(doc, "7.2 ModeB", 2)
    add_para(doc, "ModeB 面向客户案例和跨路径优化视角，继续沿用同一套 setup 校验与产能占用逻辑。ModeB 输出中涉及工作中心、客户、产品、案例的明细时，也会继承 setup 后的 Capacity_Used_Tons。")

    add_heading(doc, "7.3 Max 与 Planned", 2)
    add_table(
        doc,
        ["口径", "产能来源", "setup 换算", "报告解读"],
        [
            ["Max", "Max Capacity", "按 Max Capacity ÷ 月小时数换算", "表示理论最大承载能力下的结果。"],
            ["Planned", "Planned Capacity", "仍按 Max Capacity ÷ 月小时数换算", "表示计划产能约束下的结果，但 setup 消耗吨位口径保持稳定。"],
        ],
        [1.1, 1.7, 2.0, 1.9],
    )

    add_heading(doc, "8. 报告排序逻辑", 1)
    add_para(doc, "为了避免使用者误解同一个月、同一个工厂、同一个工作中心、同一个产品的 setup 为什么只计算一次，分配明细会把同口径内的相关行排在一起。")
    add_table(
        doc,
        ["排序层级", "字段", "目的"],
        [
            ["1", "Capacity_Basis", "先区分 Max 与 Planned。"],
            ["2", "Month", "同一月份集中展示。"],
            ["3", "Plant", "同一工厂集中展示。"],
            ["4", "WorkCenter", "同一工作中心集中展示。"],
            ["5", "Product", "同一产品集中展示，便于理解同月只触发一次 setup。"],
            ["6", "PlannerName", "同产品下按 planner 继续归类。"],
            ["7", "Source_Resource", "区分原始资源来源。"],
            ["8", "AllocationType / RouteType / Priority", "同组内按分配类型、路径类型和优先级稳定排序。"],
        ],
        [0.8, 2.0, 3.0],
    )

    add_heading(doc, "9. 报告和列含义", 1)
    add_heading(doc, "9.1 分配明细 Allocation Detail", 2)
    add_table(
        doc,
        ["列", "含义", "背后逻辑"],
        [
            ["Capacity_Basis", "产能口径，通常为 Max 或 Planned。", "决定当前行使用哪套可用产能约束。"],
            ["Month", "需求月份。", "用于匹配需求、月小时数、工作中心产能和 setup 聚合。"],
            ["Plant", "工厂。", "用于工厂维度汇总和排序。"],
            ["PlannerName", "Planner 名称。", "用于需求归属和报表追踪。"],
            ["Product", "产品。", "用于匹配 capacity/routing，并参与 setup 聚合。"],
            ["WorkCenter", "最终分配工作中心。", "优化后实际承载该产品吨位的工作中心。"],
            ["Source_Resource", "来源工作中心或原始路径资源。", "用于区分主路径与替代路径来源。"],
            ["Allocation_Tons", "分配生产吨位。", "优化模型分配给该工作中心的产品需求吨位。"],
            ["Setup_Applied", "是否触发 setup。", "当同月同工作中心同产品累计分配吨位大于 1 吨时为 True/Yes。"],
            ["Setup_Hours", "输入维护的换模小时数。", "来自 capacity 或 routing，经一致性校验后使用。"],
            ["Setup_Equivalent_Tons_By_Max", "setup 折算吨位。", "Setup_Hours × Max Capacity ÷ 月小时数。"],
            ["Capacity_Used_Tons", "实际占用产能吨位。", "Allocation_Tons + Setup_Equivalent_Tons_By_Max。"],
            ["RouteType", "路径类型。", "标识主路径或 alternative routing。"],
            ["Priority", "路径优先级。", "用于同组内排序和优化选择参考。"],
            ["Service/Unmet 字段", "服务水平或未满足需求相关字段。", "用于判断需求是否完全满足。"],
        ],
        [1.4, 2.15, 2.7],
    )

    add_heading(doc, "9.2 Capacity Summary / WorkCenter Summary", 2)
    add_table(
        doc,
        ["列", "含义", "计算逻辑"],
        [
            ["Capacity_Basis", "Max 或 Planned。", "区分产能口径。"],
            ["Month", "月份。", "汇总的时间维度。"],
            ["WorkCenter", "工作中心。", "汇总产能压力的资源维度。"],
            ["Available_Capacity_Tons", "可用产能吨位。", "来自对应口径的 capacity/routing 输入。"],
            ["Allocated_Tons", "生产分配吨位合计。", "汇总明细中的 Allocation_Tons。"],
            ["Setup_Equivalent_Tons_By_Max", "setup 折算吨位合计。", "汇总明细中被触发 setup 的等效吨位。"],
            ["Capacity_Used_Tons", "产能占用合计。", "Allocated_Tons + Setup_Equivalent_Tons_By_Max。"],
            ["Utilization", "产能利用率。", "Capacity_Used_Tons ÷ Available_Capacity_Tons。"],
            ["Overload_Tons", "超产能吨位。", "max(Capacity_Used_Tons - Available_Capacity_Tons, 0)。"],
        ],
        [1.55, 2.3, 2.7],
    )

    add_heading(doc, "9.3 ModeB Customer / Case Report", 2)
    add_table(
        doc,
        ["列", "含义", "计算逻辑"],
        [
            ["Customer/Case", "客户或案例维度。", "来自 planner load 或 ModeB 分析输入。"],
            ["Product", "产品。", "用于关联需求和路径。"],
            ["Demand_Tons", "需求吨位。", "客户或案例维度下的需求汇总。"],
            ["Allocated_Tons", "已分配吨位。", "优化结果中满足的吨位。"],
            ["Unmet_Tons", "未满足吨位。", "Demand_Tons - Allocated_Tons。"],
            ["Service_Level", "服务水平。", "Allocated_Tons ÷ Demand_Tons；需求为 0 时按报告逻辑显示为空或 100%。"],
            ["WorkCenter", "承载工作中心。", "显示该客户/案例产品最终使用的资源。"],
            ["Setup 相关列", "换模小时和折算吨位。", "沿用同月同工作中心同产品只触发一次的逻辑。"],
        ],
        [1.45, 2.15, 2.8],
    )

    add_heading(doc, "10. 常见问题", 1)
    add_table(
        doc,
        ["问题", "原因", "处理方式"],
        [
            ["报告中某个产品 setup 为 0", "该 Product + WorkCenter 的 Setup_Hours 输入为 0，或该组累计分配吨位未超过 1 吨。", "检查 capacity/routing 的 Setup_Hours 和该月明细累计吨位。"],
            ["同一个产品多行只有一行显示 setup", "工具按 Month + WorkCenter + Product 聚合，同月只触发一次 setup。", "这是预期逻辑；明细排序已把同组行排在一起便于解释。"],
            ["工具停止并生成错误报告", "输入校验失败，例如 setup 为空、非数字、负数，或 capacity/routing 数值不一致。", "按错误报告列出的 Product + Resource 修正输入。"],
            ["Planned 报告中 setup 折算看起来不是 Planned 速率", "v2.2.1 明确 setup 折算统一使用 Max Capacity。", "这是为了保证 Max 与 Planned 月份和工作中心不变时 setup 消耗保持一致。"],
            ["Word 文档显示问号", "文档生成过程发生编码损坏。", "使用本脚本重新生成，确保 Python 源文件为 UTF-8，并用 python-docx 写入。"],
        ],
        [1.7, 2.25, 2.35],
    )

    add_heading(doc, "11. 发布流程建议", 1)
    add_numbered(
        doc,
        [
            "运行 .\\scripts\\bootstrap_dev.ps1 安装运行和开发依赖。",
            "运行 .\\scripts\\release_preflight.ps1 检查 Python 包、测试、LibreOffice 和 Poppler。",
            "更新 app/version.py、packaging/*.version.txt 和 docs/CHANGELOG.md。",
            "运行 python -m pytest。",
            "运行 powershell -ExecutionPolicy Bypass -File packaging\\build_onefolder.ps1 -Target CapacityOptimizer -Clean -CreateZip。",
            "检查 dist/CapacityOptimizer/CapacityOptimizer.exe 和 delivery_packages/CapacityOptimizer-vX.Y.Z-win64.zip。",
            "打开 CapacityOptimizer.exe 或 CapacityOptimizerLauncher.pyw，确认启动器参数、运行按钮、输出目录和日志入口可用。",
            "提交代码并推送 GitHub。",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Chemical Capacity Optimizer v2.2.1 操作指引")
    set_run_font(fr, size=8, color="6B7280")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
